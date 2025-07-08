import os
import json
import pandas as pd
import numpy as np
from typing import Dict, Any, List, Optional
from docx import Document
import re
from io import BytesIO
from collections import defaultdict
import io
import openai

class DataProcessor:
    def __init__(self):
        pass
    
    async def run_statistical_test(
        self,
        file_content: bytes,
        file_name: str,
        test_type: str = "auto",
        question_key: str = ""
    ) -> Dict[str, Any]:
        """통계 분석 실행"""
        try:
            # 파일을 DataFrame으로 로드
            if file_name.endswith('.xlsx') or file_name.endswith('.xls'):
                df = pd.read_excel(io.BytesIO(file_content))
            elif file_name.endswith('.csv'):
                df = pd.read_csv(io.BytesIO(file_content))
            else:
                raise ValueError(f"지원하지 않는 파일 형식: {file_name}")
            
            # 기본 통계 계산
            numeric_columns = df.select_dtypes(include=[np.number]).columns.tolist()
            
            stats = {
                "file_name": file_name,
                "test_type": test_type,
                "question_key": question_key,
                "total_rows": len(df),
                "total_columns": len(df.columns),
                "numeric_columns": numeric_columns,
                "missing_values": df.isnull().sum().to_dict(),
                "basic_stats": {}
            }
            
            # 숫자 컬럼에 대한 기본 통계
            for col in numeric_columns:
                stats["basic_stats"][col] = {
                    "mean": float(df[col].mean()),
                    "std": float(df[col].std()),
                    "min": float(df[col].min()),
                    "max": float(df[col].max()),
                    "count": int(df[col].count())
                }
            
            # 상관관계 분석
            if len(numeric_columns) > 1:
                correlation_matrix = df[numeric_columns].corr()
                stats["correlation"] = correlation_matrix.to_dict()
            
            return {
                "success": True,
                "result": stats
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def process_survey_data(self, df: pd.DataFrame) -> Dict[str, Any]:
        """설문 데이터 처리"""
        # 설문 데이터 특화 처리 로직
        result = {
            "question_keys": [],
            "question_texts": {},
            "tables": {}
        }
        
        # 질문 키 추출 (첫 번째 컬럼에서)
        if len(df.columns) > 0:
            first_col = df.columns[0]
            question_keys = df[first_col].dropna().unique().tolist()
            result["question_keys"] = question_keys
            
            # 각 질문별 테이블 생성
            for key in question_keys:
                if key in df[first_col].values:
                    question_data = df[df[first_col] == key]
                    result["tables"][key] = {
                        "columns": question_data.columns.tolist(),
                        "data": question_data.values.tolist()
                    }
        
        return result 

    def process_excel_file(self, file_content: bytes, file_name: str):
        """엑셀/CSV 파일의 bytes를 DataFrame으로 변환"""
        if file_name.endswith('.xlsx') or file_name.endswith('.xls'):
            return pd.read_excel(io.BytesIO(file_content))
        elif file_name.endswith('.csv'):
            return pd.read_csv(io.BytesIO(file_content))
        else:
            raise ValueError(f"지원하지 않는 파일 형식: {file_name}")

    def extract_tables_from_excel(self, file_content: bytes, file_name: str = 'data.xlsx'):
        """엑셀/CSV 파일의 bytes에서 질문별 테이블 딕셔너리 반환"""
        df = self.process_excel_file(file_content, file_name)
        return self.process_survey_data(df)["tables"]

    async def analyze_table(self, table_data):
        """테이블 데이터(딕셔너리 또는 DataFrame)에 대한 간단 통계 분석 반환"""
        # table_data가 dict이면 DataFrame으로 변환
        if isinstance(table_data, dict):
            df = pd.DataFrame(table_data["data"], columns=table_data["columns"])
        else:
            df = table_data
        # 숫자 컬럼에 대한 기본 통계
        numeric_columns = df.select_dtypes(include=[float, int]).columns.tolist()
        stats = {}
        for col in numeric_columns:
            stats[col] = {
                "mean": float(df[col].mean()),
                "std": float(df[col].std()),
                "min": float(df[col].min()),
                "max": float(df[col].max()),
                "count": int(df[col].count())
            }
        return stats

    def load_survey_tables(self, file_content: bytes, file_name: str, sheet_name: str = "통계표"):
        df = pd.read_excel(io.BytesIO(file_content), sheet_name=sheet_name, header=None)

        pattern = r"^[A-Z]+\d*[-.]?\d*\."
        question_indices = df[df[0].astype(str).str.match(pattern)].index.tolist()

        tables = {}
        question_texts = {}
        question_keys = []
        key_counts = defaultdict(int)

        for i, start in enumerate(question_indices):
            end = question_indices[i + 1] if i + 1 < len(question_indices) else len(df)
            title = str(df.iloc[start, 0]).strip()

            match = re.match(pattern, title)
            if not match:
                continue
            base_key = match.group().rstrip(".")
            key_counts[base_key] += 1
            suffix = f"_{key_counts[base_key]}" if key_counts[base_key] > 1 else ""
            final_key = base_key + suffix
            final_key_norm = self.normalize_key(final_key) if hasattr(self, 'normalize_key') else final_key.replace("-", "_").replace(".", "_")

            question_texts[final_key_norm] = title + "(전체 단위 : %)"
            question_keys.append(final_key_norm)

            table = df.iloc[start + 1:end].reset_index(drop=True)

            if len(table) >= 2:
                first_header = table.iloc[0].fillna('').astype(str)
                second_header = table.iloc[1].fillna('').astype(str)

                title_text = None
                title_col_idx = None
                for idx, val in enumerate(first_header):
                    if idx > 2 and isinstance(val, str) and len(val) > 0:
                        if val not in ['관심없다', '보통', '관심있다', '평균']:
                            title_text = val
                            title_col_idx = idx
                            break

                new_columns = []
                for idx in range(len(first_header)):
                    if idx == 0:
                        new_columns.append("대분류")
                    elif idx == 1:
                        new_columns.append("소분류")
                    elif idx == 2:
                        new_columns.append("사례수")
                    else:
                        first_val = "" if (title_col_idx is not None and first_header.iloc[idx] == title_text) else first_header.iloc[idx]
                        combined = (first_val + " " + second_header.iloc[idx]).strip().replace('nan', '').strip()
                        new_columns.append(combined)

                table = table.drop([0, 1]).reset_index(drop=True)
                table.columns = new_columns
                table = table.dropna(axis=1, how='all')
                table = table.dropna(axis=0, how='all')
                table["대분류"] = table["대분류"].ffill()
                table = table.dropna(subset=["대분류", "사례수"], how="all").reset_index(drop=True)
                if len(table) > 2:
                    table = table.iloc[:-1].reset_index(drop=True)

                for col in table.columns:
                    try:
                        numeric_col = pd.to_numeric(table[col], errors='coerce')
                        if numeric_col.notna().any():
                            table[col] = numeric_col.round(1)
                    except:
                        continue

                tables[final_key_norm] = table

        return {
            "tables": tables,
            "question_texts": question_texts,
            "question_keys": question_keys
        }

    def normalize_key(self, key: str) -> str:
        return key.replace("-", "_").replace(".", "_")

    def rule_based_test_type_decision(self, columns, question_text=""):
        # 1. multi_response_keywords로 임의 분석 판단
        multi_response_keywords = [
            "1+2", "1+2+3", "복수", "다중", "multiple", "rank", "ranking", "우선순위"
        ]
        text_to_check = (" ".join(columns) + " " + question_text).lower()
        if any(keyword.lower() in text_to_check for keyword in multi_response_keywords):
            return "manual"
        # 2. categorical_patterns로 ft_test 판단
        categorical_patterns = [
            r"전혀\s*관심", r"관심\s*없(다|는)", r"관심\s*있(다|는)", r"매우\s*관심", r"관심",
            r"매우\s*만족", r"만족", r"불만족", r"매우\s*불만족", r"보통",
            r"찬성", r"반대", r"매우\s*찬성", r"매우\s*반대", r"대체로\s*찬성", r"대체로\s*반대",
            r"매우\s*중요", r"중요", r"그다지\s*중요하지\s*않", r"전혀\s*중요하지\s*않",
            r"매우\s*심각", r"심각", r"심각하지\s*않", r"전혀\s*심각하지\s*않",
            r"자주", r"가끔", r"거의\s*없", r"전혀\s*없",
            r"안전", r"매우\s*안전", r"위험", r"매우\s*위험",
            r"들어본\s*적", r"사용한\s*적", r"경험했", r"인지",
            r"의향", r"생각", r"예정", r"계획", r"할\s*것",
            r"매우", r"약간", r"보통", r"그다지", r"전혀"
        ]
        if any(any(re.search(pattern, col) for pattern in categorical_patterns) for col in columns):
            return "ft_test"
        # 3. 나머지는 chi_square
        return "chi_square"

    async def llm_test_type_decision(self, columns, question_text=""):
        # 1. multi_response_keywords로 임의 분석 판단
        multi_response_keywords = [
            "1+2", "1+2+3", "복수", "다중", "multiple", "rank", "ranking", "우선순위"
        ]
        text_to_check = (" ".join(columns) + " " + question_text).lower()
        if any(keyword.lower() in text_to_check for keyword in multi_response_keywords):
            return "manual"
        # 2. LLM 프롬프트 생성
        TEST_TYPE_PROMPT = """
당신은 통계 전문가입니다.

아래는 설문 응답 결과 테이블의 열 이름 목록입니다. 이 열들은 응답자들이 선택하거나 평가한 설문 문항의 결과로 구성된 통계표입니다.

당신의 임무는, 이 테이블이 **어떤 통계 검정(F/T-test 또는 Chi-square)** 에 적합한지를 판단하는 것입니다.

📋 열 이름 목록:
{column_names}

---
Let's think step by step

판단 기준:

- `ft_test` (연속형 수치 응답):
    - 문항이 1~5점 척도, 평균, 비율, 점수 등 숫자 기반으로 요약되어 있다면 F-test 또는 T-test가 적절합니다.
    - 예시 열 이름: "평균", "만족도 점수", "~% 비율", "5점 척도", "평균 점수", "관심도 평균"
    - "전혀 관심이 없다", "매우 관심 있다" 등은 실제로는 선택지이지만, 빈도나 비율로 수치화되었을 경우 → 연속형으로 판단

- `chi_square` (범주형 선택 응답):
    - 문항이 응답자들이 특정 항목을 **선택**하거나 **다중선택**한 결과일 경우, 범주형 응답으로 보고 카이제곱 검정이 적합합니다.
    - 예시 열 이름: "주요 이용시설", "선택 이유", "가장 많이 선택한 장소", "다중 응답"

- 'manual' (복수 응답 형 질문):
    - 문항의 내용이 복수 응답 키워드(1+2  순위, 1+2+3 순위, 복수응답 등)가 있거나 사용자가 단일 응답이 아닌 복수 응답인 경우 임의 분석 (manual)이 적합합니다.
        - **1순위**만 있을 경우, 이는 복수 응답이 아닌 단일 응답으로 간주하여 ft_test 또는 chi-square 중 하나로 판단해주세요.
    - 예시 문항: 서울의 대기환경 개선을 위해 서울시가 가장 역점을 두고 추진해야 할 분야는 무엇입니까? (1+2순위)(전체 단위 : %)

❗ 오판 주의:
- 응답 선택지 이름(예: "전혀 관심 없다", "매우 관심 있다")가 열 이름에 포함되더라도, **비율, 평균 등의 수치형 요약**이면 `ft_test`로 간주합니다.
- 테이블이 전체적으로 평균값 또는 %비율 중심이면 `ft_test` 선택이 더 적절합니다.

---

📌 답변 형식: 아래의 형식처럼 선택의 이유에 대해서 답변하지 말고 "적합한 통계 검정의 방법만" 출력하세요.

- 반드시 다음 중 하나로만 답해주세요 (소문자):
    - ft_test
    - chi_square

적합한 통계 방법: (ft_test 또는 chi_square)
"""
        prompt = TEST_TYPE_PROMPT.format(column_names=", ".join(columns))
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise RuntimeError("OPENAI_API_KEY 환경 변수가 설정되어 있지 않습니다.")
        response = await openai.AsyncOpenAI(api_key=api_key).chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "당신은 통계 전문가입니다."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.0,
            max_tokens=20
        )
        content = response.choices[0].message.content.strip().lower()
        if "chi" in content:
            return "chi_square"
        elif "ft" in content:
            return "ft_test"
        else:
            return "unknown"

def extract_guide_subjects_from_docx_bytes(docx_bytes: bytes):
    """
    DOCX 파일의 bytes를 받아서, 본문과 표에서 주제/질문(섹션명, 질문 등)을 robust하게 추출한다.
    """
    doc = Document(BytesIO(docx_bytes))
    # 1. 본문 텍스트 추출
    paragraph_texts = [para.text for para in doc.paragraphs if para.text.strip() != ""]
    # 2. 표 텍스트 추출
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_texts.append("\t".join(row_text))  # 탭으로 구분
    # 3. 전체 텍스트 합치기
    full_text = "\n".join(paragraph_texts + table_texts)
    # 4. 주제/질문 robust 추출
    subjects = []
    for line in full_text.split('\n'):
        line = line.strip()
        # 숫자+마침표/괄호/콜론, ▷, -, •, PART, 소제목 등
        if re.match(r'^(PART|[0-9]+[.)]|[0-9]+:|▷|-|•)', line):
            subjects.append(line)
        # 질문형 문장
        elif re.search(r'(무엇|어떻게|있나요|생각하시나요|알고 계시나요|이유|방식|방안|의견|경험|느낀 점|추천|평가|의미|정의|차이|특징|장점|단점|문제|해결|필요|중요|역할|기대|효과|방향|계획|전략|방법|있으신가요|있습니까|있을까요|있을지)', line):
            subjects.append(line)
    return subjects 