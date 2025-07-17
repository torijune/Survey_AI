from typing import List, Dict, Any, Optional
import re
import io
from docx import Document
from app.fgi.domain.entities import FGIState, FGIChunk, FGIAnalysisResult, FGIGuideSubject
from app.fgi.infra.openai_client import OpenAIClient


class FGIAnalysisService:
    """FGI 분석 비즈니스 로직 서비스"""
    
    def __init__(self, openai_client: OpenAIClient):
        self.openai_client = openai_client
    
    async def extract_text_from_docx(self, state: FGIState) -> FGIState:
        """DOCX 파일에서 텍스트 추출"""
        try:
            if not state.file_content:
                raise ValueError("파일 내용이 없습니다.")
            
            doc = Document(io.BytesIO(state.file_content))
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            # 표에서 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join([cell.text.strip() for cell in row.cells])
                    text += '\n' + row_text
            
            print("[FGI] DOCX 텍스트 추출 완료")
            print(f"[FGI] DOCX 텍스트 길이: {len(text)}")
            print(f"[FGI] DOCX 텍스트 앞부분: {text[:100]}")
            
            state.extracted_text = text
            return state
            
        except Exception as e:
            raise Exception(f"DOCX 텍스트 추출 실패: {str(e)}")
    
    async def extract_text_from_txt(self, state: FGIState) -> FGIState:
        """TXT 파일에서 텍스트 추출"""
        try:
            if not state.file_content:
                raise ValueError("파일 내용이 없습니다.")
            
            text = state.file_content.decode('utf-8')
            print("[FGI] TXT 텍스트 추출 완료")
            
            state.extracted_text = text
            return state
            
        except Exception as e:
            raise Exception(f"TXT 텍스트 추출 실패: {str(e)}")
    
    def split_text(self, text: str, max_length: int = 4000) -> List[str]:
        """텍스트를 청크로 분할"""
        paragraphs = text.split('\n')
        chunks = []
        current = ''
        
        for para in paragraphs:
            if len(current) + len(para) < max_length:
                current += para + '\n'
            else:
                chunks.append(current)
                current = para + '\n'
        
        if current:
            chunks.append(current)
        
        return chunks
    
    def split_long_text_with_overlap(self, text: str, max_chunk_size: int, overlap_size: int) -> List[str]:
        """오버랩을 포함한 긴 텍스트 분할"""
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = min(start + max_chunk_size, text_length)
            chunk = text[start:end]
            chunks.append(chunk)
            start += max_chunk_size - overlap_size
        
        return chunks
    
    def smart_chunk_with_overlap(self, text: str, min_chunk_size: int = 1000, 
                                max_chunk_size: int = 4000, overlap_size: int = 300) -> List[str]:
        """스마트 청킹 (Q&A 블록 기반)"""
        pattern = r'(사회자\s*:\s*)'
        blocks = re.split(pattern, text)
        qa_blocks = []
        i = 1
        
        while i < len(blocks):
            block = blocks[i] + (blocks[i+1] if i+1 < len(blocks) else "")
            qa_blocks.append(block.strip())
            i += 2
        
        chunks = []
        current = ""
        
        for block in qa_blocks:
            if len(current) + len(block) < min_chunk_size:
                current += "\n" + block
            elif len(current) + len(block) < max_chunk_size:
                current += "\n" + block
                for part in self.split_long_text_with_overlap(current.strip(), max_chunk_size, overlap_size):
                    chunks.append(part)
                current = ""
            else:
                if current:
                    for part in self.split_long_text_with_overlap(current.strip(), max_chunk_size, overlap_size):
                        chunks.append(part)
                current = block
        
        if current:
            for part in self.split_long_text_with_overlap(current.strip(), max_chunk_size, overlap_size):
                chunks.append(part)
        
        return chunks
    
    def extract_qa_blocks(self, text: str) -> List[str]:
        """Q&A 블록 추출 및 오버랩 청킹 적용"""
        return self.smart_chunk_with_overlap(text, min_chunk_size=1000, max_chunk_size=4000, overlap_size=300)
    
    def merge_chunks(self, chunks: List[str], group_size: int = 3) -> List[str]:
        """청크들을 그룹으로 병합"""
        merged = []
        for i in range(0, len(chunks), group_size):
            merged_chunk = chunks[i]
            for j in range(1, group_size):
                if i + j < len(chunks):
                    merged_chunk += '\n' + chunks[i + j]
            merged.append(merged_chunk)
        return merged
    
    async def analyze_chunk_with_llm(self, chunk: str, on_step=None) -> str:
        """LLM을 사용한 청크 분석"""
        if on_step:
            on_step('[FGI] LLM 요약 시작')
        print('[FGI] LLM 요약 시작')
        
        guide_prompt = (
            "이 청크는 회의의 일부분입니다. 사회자와 참여자들의 대화 내용을 자세히 분석해 주세요.\n\n"
            "다음 형식으로 분석해 주세요:\n"
            "1. 사회자가 한 질문/발언: [사회자의 질문이나 발언 내용]\n"
            "2. 참여자들의 답변/의견: [참여자들의 구체적인 답변과 의견]\n"
            "3. 주요 논의 포인트: [이 대화에서 나온 핵심 내용]\n"
            "4. 청크 요약: [이 청크 전체의 핵심 내용을 구체적인 예시와 함께 요약]\n\n"
            "각 참여자의 발언이 구체적으로 무엇인지, 어떤 의견을 제시했는지 자세히 분석해 주세요.\n"
            "청크 요약에서는 참여자들이 제시한 구체적인 의견이나 예시를 포함하여 작성해 주세요.\n"
        )
        
        prompt = guide_prompt + f"\n아래는 회의록 텍스트입니다.\n{chunk}\n"
        
        messages = [
            {"role": "system", "content": "당신은 회의록 분석가입니다."},
            {"role": "user", "content": prompt}
        ]
        
        summary = await self.openai_client.call(messages)
        
        if on_step:
            on_step('[FGI] LLM 요약 완료')
        print('[FGI] LLM 요약 완료')
        
        return summary
    
    async def create_final_summary_with_llm(self, all_summaries: str, on_step=None, 
                                          system_prompt: Optional[str] = None) -> str:
        """최종 요약 생성"""
        prompt = (
            "당신은 설문 조사를 위해 진행된 회의의 결과를 분석하여 참여자들의 의견을 추출하고 요약하는 전문가입니다."
            "회의에서 논의된 내용들을 참고하여 전체 회의에서의 참여자들의 의견을 요약하여 작성해주세요."
            "각 회의 청크별 분석들은 사회자가 진행한 주제에 대한 참여자들의 의견을 분석한 것입니다.\n\n"
            f"아래는 각 청크별 분석 결과입니다.\n"
            f"각 청크 요약들 속에 요약되어 있는 내용들에 대해서 주제를 직접 정하여 각 주제별로 회의에서 어떤 의견들이 나왔는지 키워드 중심의 요약과 자세한 분석을 함께 진행해주세요."
            f"키워드 중심으로 요약할 때에는 참여자들의 제시한 중심 주제와 키워드만을 제공해주세요."
            f"자세한 요약을 진행할 때에는 참여자들이 제시한 의견들에 대해서 나누었다., 제시하였다., 논의되었다. 등의 표현이 아니라 참여자들이 구체적으로 어떤 의견을 제시했는지 자세히 요약하여 분석해주세요."
            f"{all_summaries}\nLet's think step by step."
        )
        
        print('[FGI][최종요약 LLM 프롬프트]\n' + prompt)
        
        system_prompt_to_use = system_prompt or '당신은 회의록 요약가입니다.'
        
        messages = [
            {"role": "system", "content": system_prompt_to_use},
            {"role": "user", "content": prompt}
        ]
        
        summary = await self.openai_client.call(messages)
        return summary
    
    async def extract_guide_subjects_llm(self, file_content: bytes) -> List[str]:
        """LLM을 사용한 가이드 주제 추출"""
        try:
            doc = Document(io.BytesIO(file_content))
            full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            # 표에서 텍스트 추출
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join([cell.text.strip() for cell in row.cells])
                    table_texts.append(row_text)
            
            if table_texts:
                full_text += '\n' + '\n'.join(table_texts)
            
            print('[FGI] 가이드라인 전체 텍스트 앞부분:', full_text[:300])

            # LLM 프롬프트
            prompt = (
                "아래 가이드라인 문서 전체에서 실제 토론/질문/주제/파트/소주제만 리스트로 뽑아줘. "
                "서론, 결론, 설명문, 안내문 등은 절대 포함하지 마. "
                "리스트는 반드시 숫자(1., 2., ...) 또는 특수문자(-, •, ▷ 등)로 시작하는 항목만 포함해. "
                "주제 리스트 이외의 설명, 안내, 서론, 결론 등은 절대 포함하지 마."
                "주제 리스트에 대해서만 출력하고 다른 내용은 절대 포함하지마.\n\n"
                + full_text +
                "\n예시: 주제1\n 주제2\n 주제3\n 주제4\n 주제5\n  ...\n\n"
                "Let's think step by step."
            )
            
            messages = [
                {"role": "system", "content": "당신은 FGI/설문 가이드라인 분석 전문가입니다."},
                {"role": "user", "content": prompt}
            ]
            
            llm_response = await self.openai_client.call(messages)
            
            # LLM 응답을 줄 단위로 파싱 + 불필요한 줄 필터링
            filter_phrases = [
                "아래는", "주제", "이 주제들은", "기초 자료로 활용될 수 있습니다", "설명", "서론", "결론"
            ]
            
            topics = [
                line.strip('-•\t ')
                for line in llm_response.split('\n')
                if line.strip()
                and not any(s in line for s in filter_phrases)
                and (re.match(r'^([0-9]+[.)]|[0-9]+:|▷|-|•)', line.strip()) or len(line.strip()) > 5)
            ]
            
            # 앞 숫자/기호 제거
            topics = [re.sub(r'^\s*[\d]+[.)\-:•▷]?\s*', '', t) for t in topics if len(t) > 2]
            
            if topics:
                print(f"[FGI][LLM] 추출된 주제 {len(topics)}개:")
                for i, t in enumerate(topics):
                    print(f"  [{i+1}] {t}")
                return topics
            
            return []
            
        except Exception as e:
            print(f"[FGI][LLM] 주제 추출 실패, rule-based로 대체: {e}")
            return self.extract_guide_subjects_rule_based(file_content)
    
    def extract_guide_subjects_rule_based(self, file_content: bytes) -> List[str]:
        """Rule-based 가이드 주제 추출"""
        try:
            doc = Document(io.BytesIO(file_content))
            full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            lines = full_text.split('\n')
            subjects = []
            
            for line in lines:
                trimmed = line.strip()
                if re.match(r'^(PART|[0-9]+[.)]|[0-9]+:|▷|-|•)', trimmed):
                    subjects.append(trimmed)
                elif re.search(r'(무엇|어떻게|있나요|생각하시나요|알고 계시나요|이유|방식|방안|의견|경험|느낀 점|추천|평가|의미|정의|차이|특징|장점|단점|문제|해결|방안|필요|중요|역할|기대|효과|방향|계획|전략|방법|있으신가요|있습니까|있을까요|있을지)', trimmed):
                    subjects.append(trimmed)
            
            print(f"[FGI] 추출된 가이드 주제 {len(subjects)}개:")
            for i, subject in enumerate(subjects):
                print(f"  [{i+1}] {subject}")
            
            return subjects
            
        except Exception as e:
            print(f"[FGI] Rule-based 주제 추출 실패: {e}")
            return [] 